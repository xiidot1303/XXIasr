import datetime
import requests
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required, permission_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import UpdateView, CreateView
from django.contrib.auth.models import User
from django.shortcuts import redirect, render
from django.db.models import Sum
from base.forms import *
from .models import SMS, Access, Client, Notes, Profile, Service, Task, Upload, SMStext, telegramPost, subscriptions, Bot_user
from django.db.models.query import Q
from django.contrib import messages
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.views.decorators.csrf import csrf_exempt
from bot.update import dp, updater
from django.http.response import HttpResponse, FileResponse, Http404
from data.config import ENVIRONMENT
from telegram import Update
import json
import telegram
from base.utils.message import *
from base.utils import services
from base.utils.services import number_format as nf, handle_uploaded_file
import os
from control.settings import BASE_DIR
from base.templatetags.styles import car_number
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from django.core.exceptions import PermissionDenied

# Create your views here.
def loginPage(request):
    if request.user.is_authenticated:
        return redirect('home')
    else:
        if request.method == "POST":
            username = request.POST['username']
            password = request.POST['password']

            try:
                user = User.objects.get(username=username)
            except:
                print("Foydalanuvchi topilmadi")

            user = authenticate(request,username=username, password=password)

            if user is not None:
                login(request, user)
                profile = Profile.objects.get(user=user)
                return redirect("home")
            else:
                print("Login yoki Parol xato")
    return render(request, 'login.html')

def  logoutPage(request):
    logout(request)
    return redirect('login')

@login_required(login_url='login')
def homePage(request):
    # check ip address
    ip = services.get_user_ip(request)
    from data.config import ALLOWED_IPS

    if not ip in ALLOWED_IPS and not '*' in ALLOWED_IPS:
        logout(request)
        raise PermissionDenied

    profile = Profile.objects.get(user=request.user)
    if profile.status == 'admin':
        uploads  = Upload.objects.all()
        uploads_uncompleted = Upload.objects.filter(status='5')
        uploads_unpriced = Upload.objects.filter(status='0')
        tasks_completed  = Task.objects.filter(status='10')
        tasks_uncompleted = Task.objects.filter(Q(status='0') | Q(status='5'))
    elif profile.status == 'superuser':
        uploads  = Upload.objects.all()
        uploads_uncompleted = Upload.objects.filter(Q(status='5', office=profile.office) & Q(office=profile.office))
        uploads_unpriced = Upload.objects.filter(Q(status='0') & Q(office=profile.office))
        tasks_completed  = Task.objects.filter(Q(status='5', office=profile.office) & Q(user=profile))
        tasks_uncompleted = Task.objects.filter(Q(user=profile) & Q(status='0', office=profile.office) | Q(status='5', office=profile.office))
    elif profile.status == 'user':
        uploads = ""
        uploads_uncompleted = Upload.objects.filter(Q(status='5') & Q(reciever=profile))
        uploads_unpriced = Upload.objects.filter(Q(status='0') & Q(reciever=profile))
        tasks_completed  = Task.objects.filter(Q(status='5') & Q(user=profile))
        tasks_uncompleted = Task.objects.filter(Q(status='0') | Q(status='5'))
        tasks_uncompleted = tasks_uncompleted.filter(Q(user=profile))
    notes_process = Notes.objects.filter(Q(status='5') & Q(user=profile))
    notes_uncompleted = Notes.objects.filter(Q(status='0') & Q(user=profile))
    context = {'uploads':uploads, 'uploads_uncompleted':uploads_uncompleted,'uploads_unpriced':uploads_unpriced,'tasks_completed':tasks_completed,'tasks_uncompleted':tasks_uncompleted, 'profile':profile, 'notes_process':notes_process, 'notes_uncompleted':notes_uncompleted}
    return render(request, 'base/home.html', context)

@login_required(login_url='login')
def monitoringPage(request):
    profile = Profile.objects.get(user=request.user)
    users = Profile.objects.all().exclude(status = 'admin')
    services = Service.objects.all()

    if 'filter' in request.GET:
        date_from = request.GET['from']
        if not date_from:
            date_from = '1000-12-12'
        date_to = request.GET['to']
        if not date_to:
            date_to = '3000-12-12'
        office = request.GET['office']
        status = request.GET['status']
        user = request.GET['user']
        service = request.GET['service']
        
        uploads = Upload.objects.filter(period__range=(date_from, date_to))
        if office:
            uploads = uploads.filter(office=office)
        if status:
            uploads = uploads.filter(status=status)
        if user:
            uploads = uploads.filter(reciever__id = user)
        if service:
            uploads = uploads.filter(service__id = service)


    else:
        uploads = Upload.objects.all()
    if profile.status == 'admin':
        result = uploads.filter(Q(status=5) | Q(status=0))
        context = {'profile':profile, 'result':result, 'users': users, 'services': services}
        return render(request, 'base/monitoring.html', context)

    elif profile.status == 'superuser':
        result = uploads.filter(Q(status=5, office=profile.office) | Q(status=0, office=profile.office))
        context = {'profile':profile, 'result':result, 'users': users, 'services': services}
        return render(request, 'base/monitoring.html', context)


    elif profile.status == 'user':
        uncompleted = uploads.filter(Q(status=5) | Q(status=0))
        result = uncompleted.filter(reciever=profile)
        context = {'profile':profile, 'result':result, 'services': services}
        return render(request, 'base/monitoring.html', context)
    else:
        return render(request, 'error-404.html')

@login_required(login_url='login')
def uploadPage(request):
    form = UploadForm()
    profile = Profile.objects.get(user=request.user)
    if request.method == "POST":
        profile = Profile.objects.get(id=request.POST['reciever'])
        admins = Profile.objects.filter(Q(status='admin') | Q(status='superuser'))
        sender = Profile.objects.get(user=request.user)
        form = UploadForm(request.POST, request.FILES)
        try:
            client = Client.objects.get(id=request.POST['client'])
        except:
            client = ""
        if form.is_valid():
            upload = form.save(commit=False)
            upload.office = profile.office
            upload.sender = sender
            upload.save()
            shablon1 = SMStext.objects.get(id=4).text
            text1=shablon1.replace("**nom", client.name)
            rephone = client.phone1.replace(" ", "")
            rephone = rephone.replace("-","")
            rephone = rephone.replace(".","")
            rephone = rephone.replace(")","")
            rephone = rephone.replace("(","")
            if len(rephone) == 13:
                rephone = rephone
            elif len(rephone) == 9:
                rephone = '+998' + str(rephone)
            elif len(rephone) == 12 and rephone[0] == '9':
                rephone = '+' + str(rephone)
            elif len(rephone) == 0 or len(rephone) == 1:
                rephone = False
            else:
                rephone = False

            if rephone:
                numberid = rephone
            
            url = 'http://91.204.239.44/broker-api/send'
            headers = {'Content-type': 'application/json',  # Определение типа данных
                    'Accept': 'text/plain',
                    'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
            data = {
            "messages":
            [
            {
            "recipient":numberid,
            "message-id":"prime000019953",

                "sms":{

                "originator": "21ASR",
                "content": {
                "text": text1
                }
                }
                    }
                ]
            } 
            
            requests.post(url, json=data, headers=headers)

            # send message to client via telegram bot
            bot_send_message(client, text1)
         
            return redirect('upload')

    context = {'form':form, 'profile':profile}
    return render(request, 'base/upload.html', context)

@login_required(login_url='login')
def createService(request):
    form = ServiceCreation()
    profile = Profile.objects.get(user=request.user)
    services = Service.objects.all().order_by('id')
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    result = 10
    paginator = Paginator(services, result)

    try:
        services = paginator.page(page)
    except PageNotAnInteger:
        page=1
        services = paginator.page(page)
    except:
        page = paginator.num_pages
        services = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)

    if request.method == "POST":
        form = ServiceCreation(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Muvaffaqiyatli ro`yxatga olindi ')
            return redirect('create-service')
        else:
            messages.error(request, 'Qandaydir xatolik :( ')
            return redirect('create-service')
    context = {'form':form, 'services':services, 'paginator':paginator, 'page_range':page_range, 'profile':profile}
    return render(request, 'base/createservice.html', context)

@login_required(login_url='login')
def editService(request, pk):
    service = Service.objects.get(id=pk)
    services = Service.objects.all()
    profile = Profile.objects.get(user=request.user)
    form = ServiceCreation(instance=service)
    if request.method == "POST":
        form = ServiceCreation(request.POST, instance=service)
        if form.is_valid():
            form.save()
            messages.success(request, 'Muvaffaqiyatli yangilandi ')
            return redirect('create-service')
        else:
            messages.error(request, 'Qandaydir xatolik :( ')
            return redirect('create-service')
    context = {'form':form, 'services':services, 'profile':profile}
    return render(request, 'base/createservice.html', context)

@login_required(login_url='login')
def deleteService(request, pk):
    service = Service.objects.get(id=pk)
    profile = Profile.objects.get(user=request.user)
    obj = service
    if request.method == "POST":
        service.delete()
        messages.success(request, 'Muvaffaqiyatli o`chirildi ')
        return redirect("create-service")
    return render(request, 'base/delete.html', {'obj':obj, 'profile':profile})

@login_required(login_url='login')
def taskPage(request):
    form = TaskCreation()
    profile = Profile.objects.get(user=request.user)
    
    if profile.status == 'admin':
        tasks = Task.objects.all()
    elif profile.status == 'superuser':
        tasks = Task.objects.all()
    else:
        tasks = Task.objects.filter(user=profile)
     
    date = datetime.date.today()
    
    if request.method == "POST":
        form = TaskCreation(request.POST, request.FILES)
        
        reciever = Profile.objects.get(id=request.POST['user'])
        shablon = SMStext.objects.get(id=5).text
        text = shablon.replace("**nom", reciever.name)
        rephone = reciever.phone
        numberid = rephone
        url = 'http://91.204.239.44/broker-api/send'
        headers = {'Content-type': 'application/json',  # Определение типа данных
                'Accept': 'text/plain',
                'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
        data = {
        "messages":
        [
        {
        "recipient":numberid,
        "message-id":"prime000019953",

            "sms":{

            "originator": "21ASR",
            "content": {
            "text": text
            }
            }
                }
            ]
        } 
                    
        # requests.post(url, json=data, headers=headers) # Dont send sms to staffs
        if form.is_valid():
            ts = form.save(commit=False)
            ts.sender = profile.phone
            ts.save()
            messages.success(request, 'Muvaffaqiyatli yuborildi ')
            return redirect('task')
        else:
            messages.error(request, 'Qandaydir xatolik :( ')
            return redirect('task')
    
    if 'to' and 'from' in request.GET:
        date_from = request.GET['from']
        date_to = request.GET['to']
        username = request.GET['users']
        if profile.status == 'admin' or profile.status == 'superuser':
            if username != "":
                try:
                    user = Profile.objects.get(id=username)
                except:
                    user = ""
                
        else:
            user = profile
        status = request.GET['type']
        pagType = False
        
        if username == "":
            tasks = Task.objects.filter(Q(duration__range=(date_from, date_to)))
    
        elif username != "":
            tasks = Task.objects.filter(Q(duration__range=(date_from, date_to))  & Q(user=user))
            
        if status != "":
            tasks = tasks.filter(status=status)
        else:
            tasks=tasks
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1

    pagType = True

    result = 5
    paginator = Paginator(tasks, result)

    try:
        tasks = paginator.page(page)
    except PageNotAnInteger:
        page=1
        tasks = paginator.page(page)
    except:
        page = paginator.num_pages
        tasks = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    getContext = ""

    
    users = Profile.objects.all()
    context = {'form':form, 'tasks':tasks, 'date':date, 'paginator':paginator, 'page_range':page_range, 'users':users, 'pagType':pagType, 'profile':profile}    
    return render(request, 'base/task.html', context)       

@login_required(login_url='login')
def deleteTask(request, pk):
    service = Task.objects.get(id=pk)
    profile = Profile.objects.get(user=request.user)
    obj = service
    if request.method == "POST":
        service.delete()
        messages.success(request, 'Muvaffaqiyatli o`chirildi ')
        return redirect("task")
    return render(request, 'base/delete.html', {'obj':obj, 'profile':profile})

@login_required(login_url='login')
def viewTaskPage(request, pk):
    profile = Profile.objects.get(user=request.user)
    task = Task.objects.get(id=pk)
    return render(request, 'base/viewtask.html', {'task':task, 'profile':profile})

@login_required(login_url='login')
def finishTask(request, pk):
    try:
        task = Task.objects.get(id=pk)
    except:
        task = ""
    profile = Profile.objects.get(user=request.user)
    if profile == task.user:
        if request.method == "POST":
            shablon = SMStext.objects.get(id=6).text
            text = shablon.replace("**nom", task.user.name)
            print(text)
            rephone = task.sender
            numberid = rephone
            url = 'http://91.204.239.44/broker-api/send'
            headers = {'Content-type': 'application/json',  # Определение типа данных
                    'Accept': 'text/plain',
                    'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
            data = {
            "messages":
            [
            {
            "recipient":numberid,
            "message-id":"prime000019953",
        
                "sms":{
        
                "originator": "21ASR",
                "content": {
                "text": text
                }
                }
                    }
                ]
            } 
                        
            requests.post(url, json=data, headers=headers)
            task.answer = request.POST['answer']
            try:
                task.answer_file = request.FILES['file']
            except:
                task.answer_file = ""
            task.status = request.POST['status']
            task.save()
            messages.success(request, "Jarayon yakunlandi :)")
            return redirect('task')
        return render(request, 'base/finish.html', {'task':task, 'profile':profile})
    else:
        return render(request, 'error-404.html')
    

@login_required(login_url='login')
def serviceHistoryPage(request):
    profile = Profile.objects.get(user=request.user)
    if profile.status == 'admin':
        users = Profile.objects.all()
        context = {'users':users, 'view':False, 'profile':profile}
        if 'from' and 'to' in request.GET:
            date_to = request.GET['to']
            date_from = request.GET['from']
            try:
                office = request.GET['office']
            except:
                office = ""
            try:
                status = request.GET['type']
            except:
                status = ""
            try:
                user = request.GET['users']
            except:
                user = ""

            if user != "":
                uploads = Upload.objects.filter(Q(loaded_date__range=(date_from, date_to)) & Q(office__contains=office) & Q(status__contains=status) & Q(reciever_id=user))
            else:
                uploads = Upload.objects.filter(Q(loaded_date__range=(date_from, date_to)) & Q(office__contains=office) & Q(status__contains=status))
            sum = uploads.aggregate(Sum('payment'))
            context = {'users':users, 'uploads':uploads, 'view':True, 'sum':sum, 'profile':profile} 
        return render(request, 'base/servicehistory.html', context)
    else:
        return render(request, 'error-404.html')

@login_required(login_url='login')
def createClient(request):
    profile=Profile.objects.get(user=request.user)
    context = {'view':False, 'profile':profile}
    if 'type' in request.GET:
        type = request.GET['type']
        if type == 'ytt':
            form = YaTTCreation()
            if request.method == "POST":
                form = YaTTCreation(request.POST, request.FILES)
                print(form.errors)
                if form.is_valid():
                        client = form.save(commit=False)
                        client.type = 'ytt'
                        if client.sub_type == 'aylanma':
                            client.congragulate = True
                        else:
                            client.congragulate = True
                        client.save()
                        messages.success(request, 'YaTT ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')


        if type == 'yuridik':
            form = YuridikCreation()
            if request.method == "POST":
                form = YuridikCreation(request.POST, request.FILES)
                if form.is_valid():
                        client = form.save(commit=False)
                        if client.yuridik_type == 'buxgalteriya':
                            client.congragulate = True
                        else:
                            client.congragulate = False
                        client.save()

                        client.type = 'yuridik'
                        client.save()
                        messages.success(request, 'Yuridik ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')
                else:
                    messages.error(request, 'Qandaydir xatolik')

        if type == 'jismoniy':
            form = JismoniyCreation()
            if request.method == "POST":
                form = JismoniyCreation(request.POST, request.FILES)
                if form.is_valid():
                        client = form.save(commit=False)
                        client.type = 'jismoniy'
                        client.save()
                        messages.success(request, 'Jismoniy shaxs ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')

        if type == 'taxi':
            form = TaxiCreation()
            if request.method == "POST":
                form = TaxiCreation(request.POST, request.FILES)
                if form.is_valid():
                        client = form.save(commit=False)
                        client.type = 'taxi'
                        client.save()
                        messages.success(request, 'Taxi litsenziya ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')
        

        if type == 'ishonchnoma':
            form = IshonchnomaCreation()
            if request.method == "POST":
                form = IshonchnomaCreation(request.POST, request.FILES)
                if form.is_valid():
                        client = form.save(commit=False)
                        client.type = 'ishonchnoma'
                        client.save()
                        messages.success(request, 'Ishonchnoma ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')
        
        if type == 'governor':
            form = GovernorCreation()
            if request.method == "POST":
                form = GovernorCreation(request.POST, request.FILES)
                if form.is_valid():
                        client = form.save(commit=False)
                        client.type = 'governor'
                        client.save()
                        messages.success(request, 'Hokim yordamchisi ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')
        
        if type == 'tanirovka':
            form = TanirovkaCreation()
            if request.method == "POST":
                form = YaTTCreation(request.POST, request.FILES)
                if form.is_valid():
                        client = form.save(commit=False)
                        client.type = 'tanirovka'
                        client.owner = request.POST['owner']
                        client.tex_number = request.POST['tex_number']
                        client.text_series = request.POST['text_series']
                        try:
                            if request.POST['given_date'] != "":
                                g_date = datetime.datetime.strptime(request.POST['given_date'], '%d.%m.%Y')
                                client.given_date = g_date.strftime('%Y-%m-%d')
                        except:
                            client.given_date = None
                        
                        try:
                            if request.POST['expiry_date'] != "":
                                e_date = datetime.datetime.strptime(request.POST['expiry_date'], '%d.%m.%Y')
                                client.expiry_date = e_date.strftime('%Y-%m-%d')
                        except:
                            client.expiry_date = None
                    
                        try:
                            client.ruxsatnoma = request.FILES['ruxsatnoma']
                        except:
                            client.ruxsatnoma = ""
                        try:
                            client.texpas = request.FILES['texpas']
                        except:
                            client.texpas = ""
                        
                        client.save()
                        messages.success(request, f'Mijoz ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')


        if type == 'auction':
            form = AuctionCreation()
            if request.method == "POST":
                form = AuctionCreation(request.POST, request.FILES)
                if form.is_valid():
                        client = form.save(commit=False)
                        client.type = 'auction'
                        client.save()
                        messages.success(request, 'Mijoz ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')

        if type == 'auction2':
            form = Auction2Creation()
            if request.method == "POST":
                form = Auction2Creation(request.POST, request.FILES)
                if form.is_valid():
                        client = form.save(commit=False)
                        client.type = 'auction2'
                        client.save()
                        messages.success(request, 'Mijoz ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')

        if type == 'teacher':
            form = TeacherCreation()
            if request.method == "POST":
                form = TeacherCreation(request.POST, request.FILES)
                if form.is_valid():
                        client = form.save(commit=False)
                        client.type = 'teacher'
                        client.save()
                        messages.success(request, 'Mijoz ro\'yxatga olindi')
                elif 'bot_login' in form.errors.as_data():
                    messages.error(request, 'Bunday login bilan allaqachon ro\'yxatdan o\'tilgan')
                elif 'bot_user' in form.errors.as_data():
                    messages.error(request, 'Bunday bot foydalanuvchisi boshqa bir mijozga biriktirilgan')
       
        from base.utils.services import create_key
        if request.method == "POST":
            # create key
            create_key(client.pk, profile)

        context = {'view':True, 'form':form, 'profile':profile}
        
    if 'check' in request.GET:
        try:
            c_name = request.GET['name']
        except:
            c_name = ""
            
        try:
            j = request.GET['j']
        except:
            j = ""
            
        try:
            tin = request.GET['tin']
        except:
            tin = ""
            
        if c_name !="":
            checked = Client.objects.filter(name__icontains=c_name)
        else:
            checked = Client.objects.all()
        
        if j != "":
            checked = checked.filter(jshshir__contains=j)
        else:
            checked = checked
            
        if tin !="":
            checked = checked.filter(tin__contains=tin)
        else:
            checked = checked
            
        context = {'checked':True, 'clients':checked, 'profile':profile}

        
        
    return render(request, 'base/createclient.html', context)

@login_required(login_url='login')
def YaTTPage(request):
    profile = Profile.objects.get(user=request.user)
    clients = Client.objects.filter(type='ytt').order_by('id')
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True

    result = 20
    try:
        paginator = Paginator(clients, result)
    except:
        paginator = Paginator(clients, result)
    # try:
    #     clients = paginator.page(page)
    # except PageNotAnInteger:
    #     page=1
    #     clients = paginator.page(page)
    # except:
    #     page = paginator.num_pages
    #     clients = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    key_access = Access.objects.get(name="key")
    if profile in key_access.user.all():
        access = True
    else:
        access = False
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'access':access, 'profile':profile}

    if 'filter' in request.GET:
        try:
            name = request.GET['name']
        except:
            name = ""
        try:
            stir = request.GET['stir']
        except:
            stir = ""
        try:
            tin = request.GET['tin']
        except:
            tin = ""
        try:
            cert = request.GET['certificate']
        except:
            cert = ""
        try:
            phone = request.GET['jshshir']
        except:
            phone = ""

        try:
            address = request.GET['report']
        except:
            address =""
        try:
            key = request.GET['key']
        except:
            key = ""
        try:
            key_expiry = request.GET['key_expiry']
        except:
            key_expiry = ""
        
        pagType = False

        tins = False
        
        if name != "":
            query = Client.objects.filter(Q(name__icontains=name)&Q(type='ytt'))
        else:
            query = Client.objects.filter(type='ytt')

        if stir != "":
            query = query.filter(tin__icontains=stir)

        if tin != "":
            if tin == "True":
                query = query.exclude(tin__exact="")
            else:
                query = query.filter(tin__exact="")
        else:
            query = query

        if cert != "":
            if cert == "none":
                query = query.filter(Q(guvohnoma_file__exact="") | Q(guvohnoma_file__contains=" ") | Q(guvohnoma_file__isnull=True))
            elif cert == "in_ten_days":
                today = datetime.date.today()+datetime.timedelta(days=1)
                ten_days = today+datetime.timedelta(days=10)
                query = query.filter(guvohnoma_exp__range=(today, ten_days))
            elif cert == "in_a_month":
                today = datetime.date.today()+datetime.timedelta(days=11)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(guvohnoma_exp__range=(today, ten_days))
            elif cert == "active":
                today = datetime.date.today()+datetime.timedelta(days=10)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(guvohnoma_exp__gt=ten_days)
            elif cert =="inactive":
                today = datetime.date.today()
                query = query.filter(guvohnoma_exp__lte=today)
                query = query.exclude(Q(guvohnoma_exp__isnull=True))
                
            else:
                query = query
        else:
            query = query
        
        if phone != "":
            query = query.filter(jshshir__icontains=phone)
        else:
            query = query
        

        if address != "":
            query = query.filter(sub_type=address)
        else:
            query = query

        
        if key != "":
            if key == "true":
                query = query.exclude(key__exact="")
            if key =="false":
                query = query.filter(key__exact="")
            else:
                query = query
        else:
            query=query


        if key_expiry != "":
            if key_expiry == "none":
                query = query.filter(Q(key__exact="") | Q(key__contains=" ") | Q(key__isnull=True))
            elif key_expiry == "in_ten_days":
                today = datetime.date.today()+datetime.timedelta(days=1)
                ten_days = today+datetime.timedelta(days=10)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "in_a_month":
                today = datetime.date.today()+datetime.timedelta(days=11)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "active":
                today = datetime.date.today()+datetime.timedelta(days=10)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__gt=ten_days)
            elif key_expiry =="inactive":
                today = datetime.date.today()
                query = query.filter(key_exp__lte=today)
                
            else:
                query = query
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                rephone = reciever.phone1.replace(" ", "")
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = rephone
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    } 
                    try:
                        requests.post(url, json=data, headers=headers)
                    except:
                        messages.error(request, 'Internet bilan bog\'liq muammo :(')
                else:
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
  
      
  
        context = {'clients':query, 'pagType':pagType, 'access':access, 'profile':profile}

    
    return render(request, 'base/yatt.html', context)

@login_required(login_url='login')
def YuridikPage(request):
    profile = Profile.objects.get(user = request.user)
    clients = Client.objects.filter(type='yuridik').order_by('id')
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True

    result = 20
    paginator = Paginator(clients, result)

    # try:
    #     clients = paginator.page(page)
    # except PageNotAnInteger:
    #     page=1
    #     clients = paginator.page(page)
    # except:
    #     page = paginator.num_pages
    #     clients = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    key_access = Access.objects.get(name="key")
    if profile in key_access.user.all():
        access = True
    else:
        access = False
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'access':access, 'profile':profile}

    if 'filter' in request.GET:
        try:
            name = request.GET['name']
        except:
            name = ""
        try:
            type = request.GET['type']
        except:
            type = ""
        try:
            stir = request.GET['stir']
        except:
            stir = ""
        try:
            tin = request.GET['tin']
        except:
            tin = ""
        try:
            cert = request.GET['certificate']
        except:
            cert = ""
        try:
            phone = request.GET['phone']
        except:
            phone = ""

        try:
            passport = request.GET['passport']
        except:
            passport =""
        try:
            key = request.GET['key']
        except:
            key = ""
        try:
            key_expiry = request.GET['key_expiry']
        except:
            key_expiry = ""
        
        pagType = False

        tins = False
        query = Client.objects.filter(type__exact='yuridik')
        
        if name != "":
            query = Client.objects.filter(Q(type__exact='yuridik') & Q(name__icontains=name))

        else:
            query = Client.objects.filter(type__exact='yuridik')
        
        if type != "":
            if type == 'None':
                query = query.filter(yuridik_type=None)
            else:
                query = query.filter(yuridik_type=type)

        if stir != "":
            query = query.filter(tin__icontains=stir)

        if tin != "":
            query = query.filter(jshshir__icontains=tin)

        if cert != "":
            if cert == "true":
                query = query.exclude(guvohnoma_file__isnull=False)
            else:
                query = query.filter(guvohnoma_file__exact="")
        else:
            query = query
        
        if phone != "":
            if phone == 'true':
                query = query.exclude(phone1__exact="")
            elif phone == 'false':
                query = query.filter(phone1__exact="")
            else:
                query = query
        else:
            query = query
        

        if passport != "":
            if passport == 'true':
                query = query.exclude(passport__exact="")
            if passport == 'false':
                query = query.filter(passport__exact="")
            else:
                query = query
        else:
            query = query

        
        if key != "":
            if key == "true":
                query = query.exclude(key__exact="")
            if key =="false":
                query = query.filter(key__exact="")
            else:
                query = query
        else:
            query=query


        if key_expiry != "":
            if key_expiry == "none":
                query = query.filter(Q(key__exact="") | Q(key__contains=" ") | Q(key__isnull=True))
            elif key_expiry == "in_ten_days":
                today = datetime.date.today()+datetime.timedelta(days=1)
                ten_days = today+datetime.timedelta(days=10)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "in_a_month":
                today = datetime.date.today()+datetime.timedelta(days=11)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "active":
                today = datetime.date.today()+datetime.timedelta(days=10)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__gt=ten_days)
            elif key_expiry =="inactive":
                today = datetime.date.today()
                query = query.filter(key_exp__lte=today)
                
            else:
                query = query
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                rephone = reciever.phone1.replace(" ", "")
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = str(rephone)
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    } 
                    try:
                        requests.post(url, json=data, headers=headers)
                    except:
                        messages.error(request, 'Internet bilan bog\'liq muammo :(')
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
        

        context = {'clients':query, 'pagType':pagType, 'profile':profile, 'access':access}

  
    return render(request, 'base/yuridik.html', context)

@login_required(login_url='login')
def JismoniyPage(request):
    clients = Client.objects.filter(type='jismoniy').order_by('id')
    profile = Profile.objects.get(user=request.user)
    key_access = Access.objects.get(name="key")
    if profile in key_access.user.all():
        access = True
    else:
        access = False
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True

    result = 20
    paginator = Paginator(clients, result)

    # try:
    #     clients = paginator.page(page)
    # except PageNotAnInteger:
    #     page=1
    #     clients = paginator.page(page)
    # except:
    #     page = paginator.num_pages
    #     clients = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'profile':profile}

    if 'filter' in request.GET:
        try:
            name = request.GET['name']
        except:
            name = ""
        try:
            stir = request.GET['stir']
        except:
            stir = ""
        try:
            tin = request.GET['tin']
        except:
            tin = ""
        try:
            cert = request.GET['certificate']
        except:
            cert = ""
        try:
            phone = request.GET['phone']
        except:
            phone = ""

        try:
            passport = request.GET['passport']
        except:
            passport =""
        try:
            key = request.GET['key']
        except:
            key = ""
        try:
            key_expiry = request.GET['key_expiry']
        except:
            key_expiry = ""
        
        pagType = False

        tins = False
        query = Client.objects.filter(type__exact='jismoniy')
        
        if name != "":
            query = Client.objects.filter(Q(type__exact='jismoniy') & Q(name__icontains=name))
       
        else:
            query = Client.objects.filter(type__exact='jismoniy')

        if stir != "":
            query = query.filter(tin__icontains=stir)

        if tin != "":
            query = query.filter(jshshir__icontains=tin)

        if cert != "":
            if cert == "True":
                query = query.exclude(guvohnoma_file__exact="")
            else:
                query = query.filter(guvohnoma_file__exact="")
        else:
            query = query
        
        if phone != "":
            if phone == 'true':
                query = query.exclude(phone1__exact="")
            elif phone == 'false':
                query = query.filter(phone1__exact="")
            else:
                query = query
        else:
            query = query
        

        if passport != "":
            if passport == 'true':
                query = query.exclude(passport__exact="")
            if passport == 'false':
                query = query.filter(passport__exact="")
            else:
                query = query
        else:
            query = query

        
        if key != "":
            if key == "true":
                query = query.exclude(key__exact="")
            if key =="false":
                query = query.filter(key__exact="")
            else:
                query = query
        else:
            query=query


        if key_expiry != "":
            if key_expiry == "none":
                query = query.filter(Q(key__exact="") | Q(key__contains=" ") | Q(key__isnull=True))
            elif key_expiry == "in_ten_days":
                today = datetime.date.today()+datetime.timedelta(days=1)
                ten_days = today+datetime.timedelta(days=10)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "in_a_month":
                today = datetime.date.today()+datetime.timedelta(days=11)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "active":
                today = datetime.date.today()+datetime.timedelta(days=10)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__gt=ten_days)
            elif key_expiry =="inactive":
                today = datetime.date.today()
                query = query.filter(key_exp__lte=today)
                
            else:
                query = query
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                rephone = reciever.phone1.replace(" ", "")
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = rephone
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    } 
                    try:
                        requests.post(url, json=data, headers=headers)
                    except:
                        messages.error(request, 'Internet bilan bog\'liq muammo :(')
                else:
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
  
        context = {'clients':query, 'pagType':pagType, 'access':access, 'profile':profile}

  
    return render(request, 'base/jismoniy.html', context)


@login_required(login_url='login')
def TaxiPage(request):
    clients = Client.objects.filter(type='taxi').order_by('id')
    profile = Profile.objects.get(user=request.user)
    key_access = Access.objects.get(name="key")
    if profile in key_access.user.all():
        access = True
    else:
        access = False
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True

    result = 20
    paginator = Paginator(clients, result)

    # try:
    #     clients = paginator.page(page)
    # except PageNotAnInteger:
    #     page=1
    #     clients = paginator.page(page)
    # except:
    #     page = paginator.num_pages
    #     clients = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'profile':profile}

    if 'filter' in request.GET:
        try:
            name = request.GET['name']
        except:
            name = ""
        try:
            stir = request.GET['stir']
        except:
            stir = ""
        try:
            tin = request.GET['tin']
        except:
            tin = ""
        try:
            cert = request.GET['certificate']
        except:
            cert = ""
        try:
            phone = request.GET['phone']
        except:
            phone = ""

        try:
            passport = request.GET['passport']
        except:
            passport =""
        try:
            key = request.GET['key']
        except:
            key = ""
        try:
            key_expiry = request.GET['key_expiry']
        except:
            key_expiry = ""
        
        pagType = False

        tins = False
        query = Client.objects.filter(type__exact='taxi')
        
        if name != "":
            query = Client.objects.filter(Q(type__exact='taxi') & Q(name__icontains=name))
       
        else:
            query = Client.objects.filter(type__exact='taxi')

        if stir != "":
            query = query.filter(tin__icontains=stir)

        if tin != "":
            query = query.filter(jshshir__icontains=tin)

        if cert != "":
            if cert == "True":
                query = query.exclude(guvohnoma_file__exact="")
            else:
                query = query.filter(guvohnoma_file__exact="")
        else:
            query = query
        
        if phone != "":
            if phone == 'true':
                query = query.exclude(phone1__exact="")
            elif phone == 'false':
                query = query.filter(phone1__exact="")
            else:
                query = query
        else:
            query = query
        

        if passport != "":
            if passport == 'true':
                query = query.exclude(passport__exact="")
            if passport == 'false':
                query = query.filter(passport__exact="")
            else:
                query = query
        else:
            query = query

        
        if key != "":
            if key == "true":
                query = query.exclude(key__exact="")
            if key =="false":
                query = query.filter(key__exact="")
            else:
                query = query
        else:
            query=query


        if key_expiry != "":
            if key_expiry == "none":
                query = query.filter(Q(key__exact="") | Q(key__contains=" ") | Q(key__isnull=True))
            elif key_expiry == "in_ten_days":
                today = datetime.date.today()+datetime.timedelta(days=1)
                ten_days = today+datetime.timedelta(days=10)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "in_a_month":
                today = datetime.date.today()+datetime.timedelta(days=11)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "active":
                today = datetime.date.today()+datetime.timedelta(days=10)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__gt=ten_days)
            elif key_expiry =="inactive":
                today = datetime.date.today()
                query = query.filter(key_exp__lte=today)
                
            else:
                query = query
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                rephone = reciever.phone1.replace(" ", "")
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = rephone
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    } 
                    try:
                        requests.post(url, json=data, headers=headers)
                    except:
                        messages.error(request, 'Internet bilan bog\'liq muammo :(')
                else:
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
  
        context = {'clients':query, 'pagType':pagType, 'access':access, 'profile':profile}

  
    return render(request, 'base/taxi.html', context)



@login_required(login_url='login')
def IshonchnomaPage(request):
    clients = Client.objects.filter(type='ishonchnoma').order_by('id')
    profile = Profile.objects.get(user=request.user)
    key_access = Access.objects.get(name="key")
    if profile in key_access.user.all():
        access = True
    else:
        access = False
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True

    result = 20
    paginator = Paginator(clients, result)

    # try:
    #     clients = paginator.page(page)
    # except PageNotAnInteger:
    #     page=1
    #     clients = paginator.page(page)
    # except:
    #     page = paginator.num_pages
    #     clients = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'profile':profile}

    if 'filter' in request.GET:

        try:
            name = request.GET['name']
        except:
            name = ""
        try:
            tin = request.GET['tin']
        except:
            tin = ""
        try:
            cert = request.GET['certificate']
        except:
            cert = ""
        try:
            phone = request.GET['phone']
        except:
            phone = ""

        try:
            passport = request.GET['passport']
        except:
            passport =""
        try:
            key = request.GET['key']
        except:
            key = ""
        try:
            key_expiry = request.GET['key_expiry']
        except:
            key_expiry = ""
        
        pagType = False

        tins = False
        query = Client.objects.filter(type__exact='ishonchnoma')
        
        if name != "":
            query = Client.objects.filter(Q(type__exact='ishonchnoma') & Q(name__icontains=name))
       
        else:
            query = Client.objects.filter(type__exact='ishonchnoma')


        if tin != "":
            query = query.filter(jshshir__icontains=tin)

        if cert != "":
            if cert == "True":
                query = query.exclude(guvohnoma_file__exact="")
            else:
                query = query.filter(guvohnoma_file__exact="")
        else:
            query = query
        
        if phone != "":
            if phone == 'true':
                query = query.exclude(phone1__exact="")
            elif phone == 'false':
                query = query.filter(phone1__exact="")
            else:
                query = query
        else:
            query = query
        

        if passport != "":
            if passport == 'true':
                query = query.exclude(passport__exact="")
            if passport == 'false':
                query = query.filter(passport__exact="")
            else:
                query = query
        else:
            query = query

        
        if key != "":
            if key == "true":
                query = query.exclude(key__exact="")
            if key =="false":
                query = query.filter(key__exact="")
            else:
                query = query
        else:
            query=query


        if key_expiry != "":
            if key_expiry == "none":
                query = query.filter(Q(key__exact="") | Q(key__contains=" ") | Q(key__isnull=True))
            elif key_expiry == "in_ten_days":
                today = datetime.date.today()+datetime.timedelta(days=1)
                ten_days = today+datetime.timedelta(days=10)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "in_a_month":
                today = datetime.date.today()+datetime.timedelta(days=11)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "active":
                today = datetime.date.today()+datetime.timedelta(days=10)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__gt=ten_days)
            elif key_expiry =="inactive":
                today = datetime.date.today()
                query = query.filter(key_exp__lte=today)
                
            else:
                query = query
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                rephone = reciever.phone1.replace(" ", "")
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = rephone
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    } 
                    try:
                        requests.post(url, json=data, headers=headers)
                    except:
                        messages.error(request, 'Internet bilan bog\'liq muammo :(')
                else:
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
  
        context = {'clients':query, 'pagType':pagType, 'access':access, 'profile':profile}

  
    return render(request, 'base/ishonchnoma.html', context)



@login_required(login_url='login')
def GovernorPage(request):
    clients = Client.objects.filter(type='governor').order_by('id')
    profile = Profile.objects.get(user=request.user)
    key_access = Access.objects.get(name="key")
    if profile in key_access.user.all():
        access = True
    else:
        access = False
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True

    result = 20
    paginator = Paginator(clients, result)

    # try:
    #     clients = paginator.page(page)
    # except PageNotAnInteger:
    #     page=1
    #     clients = paginator.page(page)
    # except:
    #     page = paginator.num_pages
    #     clients = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'profile':profile}

    if 'filter' in request.GET:
        try:
            name = request.GET['name']
        except:
            name = ""
        try:
            stir = request.GET['stir']
        except:
            stir = ""
        try:
            tin = request.GET['tin']
        except:
            tin = ""
        try:
            cert = request.GET['certificate']
        except:
            cert = ""
        try:
            phone = request.GET['phone']
        except:
            phone = ""

        try:
            passport = request.GET['passport']
        except:
            passport =""
        try:
            key = request.GET['key']
        except:
            key = ""
        try:
            key_expiry = request.GET['key_expiry']
        except:
            key_expiry = ""
        try:
            workplace = request.GET['workplace']
        except:
            workplace = ""
        
        pagType = False

        tins = False
        query = Client.objects.filter(type__exact='governor')
        
        if name != "":
            query = Client.objects.filter(Q(type__exact='governor') & Q(name__icontains=name))
       
        else:
            query = Client.objects.filter(type__exact='governor')

        if stir != "":
            query = query.filter(tin__icontains=stir)

        if tin != "":
            query = query.filter(jshshir__icontains=tin)

        if cert != "":
            if cert == "True":
                query = query.exclude(guvohnoma_file__exact="")
            else:
                query = query.filter(guvohnoma_file__exact="")
        else:
            query = query
        
        if phone != "":
            if phone == 'true':
                query = query.exclude(phone1__exact="")
            elif phone == 'false':
                query = query.filter(phone1__exact="")
            else:
                query = query
        else:
            query = query
        

        if passport != "":
            if passport == 'true':
                query = query.exclude(passport__exact="")
            if passport == 'false':
                query = query.filter(passport__exact="")
            else:
                query = query
        else:
            query = query

        
        if key != "":
            if key == "true":
                query = query.exclude(key__exact="")
            if key =="false":
                query = query.filter(key__exact="")
            else:
                query = query
        else:
            query=query

        
        if workplace != "":
            if workplace == "true":
                query = query.exclude(workplace__exact="")
            if workplace =="false":
                query = query.filter(workplace__exact="")
            else:
                query = query
        else:
            query=query


        if key_expiry != "":
            if key_expiry == "none":
                query = query.filter(Q(key__exact="") | Q(key__contains=" ") | Q(key__isnull=True))
            elif key_expiry == "in_ten_days":
                today = datetime.date.today()+datetime.timedelta(days=1)
                ten_days = today+datetime.timedelta(days=10)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "in_a_month":
                today = datetime.date.today()+datetime.timedelta(days=11)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__range=(today, ten_days))
            elif key_expiry == "active":
                today = datetime.date.today()+datetime.timedelta(days=10)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(key_exp__gt=ten_days)
            elif key_expiry =="inactive":
                today = datetime.date.today()
                query = query.filter(key_exp__lte=today)
                
            else:
                query = query
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                rephone = reciever.phone1.replace(" ", "")
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = rephone
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    } 
                    try:
                        requests.post(url, json=data, headers=headers)
                    except:
                        messages.error(request, 'Internet bilan bog\'liq muammo :(')
                else:
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
  
        context = {'clients':query, 'pagType':pagType, 'access':access, 'profile':profile}

  
    return render(request, 'base/governor.html', context)

@login_required(login_url='login')
def tanirovkaPage(request):
    profile = Profile.objects.get(user=request.user)
    clients = Client.objects.filter(type='tanirovka').order_by('id')
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True
    result = 20
    paginator = Paginator(clients, result)
    # try:
    #     clients = paginator.page(page)
    # except PageNotAnInteger:
    #     page=1
    #     clients = paginator.page(page)
    # except:
    #     page = paginator.num_pages
    #     clients = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'profile':profile}
    if 'filter' in request.GET:
        try:
            name = request.GET['name']
        except:
            name = ""
        try:
            owner = request.GET['owner']
        except:
            owner = ""
        try:
            tin = request.GET['tin']
        except:
            tin = ""
        try:
            tex_number = request.GET['tex_number']
        except:
            tex_number = ""
        try:
            phone = request.GET['phone']
        except:
            phone = ""
        try:
            access = request.GET['access']
        except:
            access = ""
        pagType = False
        tins = False
        query = Client.objects.filter(type__exact='tanirovka')
        if name != "":
            query = Client.objects.filter(Q(type__exact='tanirovka') & Q(name__icontains=name))
       
        else:
            query = Client.objects.filter(type__exact='tanirovka')
        if owner != "":
            query = query.filter(owner__icontains=owner)
        else:
            query = query
            
        if tin != "":
            query = query.filter(jshshir__icontains=tin)

        if tex_number != "":
            query = query.exclude(tex_number__icontains="")
        else:
            query = query

       
        
        if phone != "":
            if phone == 'true':
                query = query.exclude(phone1__exact="")
            elif phone == 'false':
                query = query.filter(phone1__exact="")
            else:
                query = query
        else:
            query = query
        



        if access != "":
            if access == "none":
                query = query.exclude(expiry_date__contains="")
            elif access == "in_ten_days":
                today = datetime.date.today()+datetime.timedelta(days=1)
                ten_days = today+datetime.timedelta(days=10)
                query = query.filter(expiry_date__range=(today, ten_days))
            elif access == "in_a_month":
                today = datetime.date.today()+datetime.timedelta(days=11)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(expiry_date__range=(today, ten_days))
            elif access == "active":
                today = datetime.date.today()+datetime.timedelta(days=10)
                ten_days = today+datetime.timedelta(days=20)
                query = query.filter(expiry_date__gt=ten_days)
            elif access =="inactive":
                today = datetime.date.today()
                query = query.filter(expiry_date__lte=today)
                
            else:
                query = query
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                rephone = reciever.phone1.replace(" ", "")
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = rephone
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    } 
                    try:
                        requests.post(url, json=data, headers=headers)
                    except:
                        messages.error(request, 'Internet bilan bog\'liq muammo :(')
                else:
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
  
        context = {'clients':query, 'pagType':pagType, 'profile':profile}

  
    return render(request, 'base/tanirovka.html', context)

@login_required(login_url='login')
def auctionPage(request):
    profile = Profile.objects.get(user=request.user)
    clients = Client.objects.filter(type='auction').order_by('id')
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True

    result = 20
    paginator = Paginator(clients, result)

    # try:
    #     clients = paginator.page(page)
    # except PageNotAnInteger:
    #     page=1
    #     clients = paginator.page(page)
    # except:
    #     page = paginator.num_pages
    #     clients = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'profile':profile}

    if 'filter' in request.GET:
        try:
            name = request.GET['name']
        except:
            name = ""
        
        try:
            owner = request.GET['owner']
        except:
            owner = ""
            
        try:
            tin = request.GET['tin']
        except:
            tin = ""
        
    
        try:
            phone = request.GET['phone']
        except:
            phone = ""

        try:
            passport = request.GET['passport']
        except:
            passport = ""
        
        
        pagType = False

        tins = False
        query = Client.objects.filter(type__exact='auction')
        
        if name != "":
            query = Client.objects.filter(Q(type__exact='auction') & Q(name__icontains=name))
         
        else:
            query = Client.objects.filter(type__exact='auction')

        if owner != "":
            query = query.filter(owner__icontains=owner)
        else:
            query = query
    
        if tin != "":
            query = query.filter(jshshir__icontains=tin)
       
        
        if phone != "":
            if phone == 'true':
                query = query.exclude(phone1__exact="")
            elif phone == 'false':
                query = query.filter(phone1__exact="")
            else:
                query = query
        else:
            query = query
        



        if passport != "":
            if passport == 'true':
                query = query.exclude(passport__exact="")
            elif passport == 'false':
                query = query.filter(passport__exact="")
            else:
                query = query
        else:
            query = query
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                rephone = reciever.phone1.replace(" ", "")
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = rephone
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    }
                    requests.post(url, json=data, headers=headers)
                    
                else:
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
  
        context = {'clients':query, 'pagType':pagType, 'profile':profile}

    return render(request, 'base/auction.html', context)

@login_required(login_url='login')
def auction2Page(request):
    profile = Profile.objects.get(user=request.user)
    clients = Client.objects.filter(type='auction2').order_by('id')
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True

    result = 20
    paginator = Paginator(clients, result)

    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'profile':profile}

    if 'filter' in request.GET:
        try:
            name = request.GET['name']
        except:
            name = ""
            
        try:
            tin = request.GET['tin']
        except:
            tin = ""
        
    
        try:
            phone = request.GET['phone']
        except:
            phone = ""

        try:
            passport = request.GET['passport']
        except:
            passport = ""
        
        
        pagType = False

        tins = False
        query = Client.objects.filter(type__exact='auction2')
        
        if name != "":
            query = Client.objects.filter(Q(type__exact='auction2') & Q(name__icontains=name))
         
        else:
            query = Client.objects.filter(type__exact='auction2')
    
        if tin != "":
            query = query.filter(jshshir__icontains=tin)
       
        
        if phone != "":
            if phone == 'true':
                query = query.exclude(phone1__exact="")
            elif phone == 'false':
                query = query.filter(phone1__exact="")
            else:
                query = query
        else:
            query = query
        



        if passport != "":
            if passport == 'true':
                query = query.exclude(passport__exact="")
            elif passport == 'false':
                query = query.filter(passport__exact="")
            else:
                query = query
        else:
            query = query
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                rephone = reciever.phone1.replace(" ", "")
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = rephone
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    }
                    requests.post(url, json=data, headers=headers)
                    
                else:
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
  
        context = {'clients':query, 'pagType':pagType, 'profile':profile}

    return render(request, 'base/auction2.html', context)

@login_required(login_url='login')
def teachersPage(request):
    profile = Profile.objects.get(user=request.user)
    clients = Client.objects.filter(type='teacher').order_by('id')
    if "page" in request.GET:
        page = request.GET['page']
    else:
        page = 1
    pagType = True

    result = 20
    paginator = Paginator(clients, result)

    # try:
    #     clients = paginator.page(page)
    # except PageNotAnInteger:
    #     page=1
    #     clients = paginator.page(page)
    # except:
    #     page = paginator.num_pages
    #     clients = paginator.page(page)
    leftIndex = int(page)-1
    if leftIndex < 1:
        leftIndex = 1
    rightIndex = (int(page)+2)
    if rightIndex > paginator.num_pages:
        rightIndex = paginator.num_pages+1
    page_range = range(leftIndex, rightIndex)
    context = {'clients':clients, 'page_range':page_range, 'paginator':paginator, 'pagType':pagType, 'profile':profile}

    if 'filter' in request.GET:
        try:
            name = request.GET['name']
        except:
            name = ""
        
        try:
            phone = request.GET['phone']
        except:
            phone = ""
        
        try:
            school = request.GET['school']
        except:
            school = ""
            
        try:
            work_as = request.GET['work_as']
        except:
            work_as = ""
            
        pagType = False

        tins = False
        
        query = Client.objects.filter(type__exact='teacher')
        
        if name != "":
            query = Client.objects.filter(Q(type__exact='teacher') & Q(name__icontains=name))
        
        else:
            query = Client.objects.filter(type__exact='teacher')
            
        if school != "":
            query = query.filter(school__icontains=school)
        else:
            query = query

        if work_as != "":
            query = query.filter(work_as__icontains=work_as)
        else:
            query = query       
        
        if phone != "":
            if phone == 'true':
                query = query.exclude(Q(phone1__exact="")& Q(phone1=None))
            elif phone == 'false':
                query = query.filter(Q(phone1__exact="") & Q(phone1=None))
            else:
                query = query
        else:
            query = query     
        
        if request.method == "POST":
            text = request.POST['text']
            for reciever in query:
                try:
                    rephone = reciever.phone1.replace(" ", "")
                except:
                    rephone = ""
                rephone = rephone.replace("-","")
                rephone = rephone.replace(".","")
                rephone = rephone.replace(")","")
                rephone = rephone.replace("(","")
                if len(rephone) == 13:
                    rephone = rephone
                    sms_status = 10
                elif len(rephone) == 9:
                    rephone = '+998' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 12 and rephone[0] == '9':
                    rephone = '+' + str(rephone)
                    sms_status = 10
                elif len(rephone) == 0 or len(rephone) == 1:
                    rephone = False
                    sms_status = 0
                else:
                    rephone = False
                    sms_status = 5

                if rephone:
                    numberid = rephone

                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )

                    url = 'http://91.204.239.44/broker-api/send'
                    headers = {'Content-type': 'application/json',  # Определение типа данных
                            'Accept': 'text/plain',
                            'Authorization': 'Basic eHhpYXNyOmJwOWJFTVA3ODI='}
                    data = {
                    "messages":
                    [
                    {
                    "recipient":numberid,
                    "message-id":"prime000019953",

                        "sms":{

                        "originator": "21ASR",
                        "content": {
                        "text": text
                        }
                        }
                            }
                        ]
                    } 
                    try:
                        requests.post(url, json=data, headers=headers)
                    except:
                        messages.error(request, 'Internet bilan bog\'liq muammo :(')
                else:
                    SMS.objects.create(
                        client=reciever,
                        text=text,
                        status=sms_status
                    )
                
  
        context = {'clients':query, 'pagType':pagType, 'profile':profile}

  
    return render(request, 'base/teacher.html', context)


@login_required(login_url='home')
def SMSPage(request):
    profile = Profile.objects.get(user = request.user)
    if profile.status == 'admin' or profile.status == 'superuser':
        sms = SMStext.objects.all()
        return render(request, 'base/smstemp.html', {'sms':sms, 'profile':profile})
    else:
        return render(request, 'error-404.html')

@login_required(login_url='home')
def editSMS(request, pk):
    profile = Profile.objects.get(user = request.user)
    if profile.status == 'admin' or profile.status == 'superuser':
        sms = SMStext.objects.get(id=pk)

        if sms:
            if request.method == "POST":
                text = request.POST['text']
                sms.text = text
                sms.save()
                messages.success(request, 'Malumot yangilandi :)')
                return redirect('sms')
            else:
                messages.error(request, 'Qandaydir xatolik :( ')
            return render(request, 'base/forms.html', {'sms':sms, 'view':'sms', 'profile':profile})
        else:
            return render(request, 'error-404.html')
    else:
        return render(request, 'error-404.html')
        
@login_required(login_url='home')
def deleteClient(request, pk):
    profile = Profile.objects.get(user = request.user)
    if profile.status == 'admin':
        if request.method == "POST":
            try:
                delclient = Client.objects.get(id=request.POST['client'])
            except:
                messages.error(request, 'Mijoz topilmadi :(')
                return redirect(request.META.get('HTTP_REFERER'))
            delclient.delete()
            messages.success(request, "Ma'lumot o'chirildi ;(")
            return redirect('create-client')
        return render(request, 'base/delete.html', {'pk':pk})
    else:
        return render(request, 'error-404.html')

@login_required(login_url='home')
def telegramPage(request):
    profile = Profile.objects.get(user = request.user)
    posts = telegramPost.objects.all().order_by('-id')
    if profile.status == 'admin' or profile.status == 'superuser':
        teleusers = subscriptions.objects.all().order_by('-id')
        clients = Client.objects.all()
        bot_users = Bot_user.objects.all()
        if request.method == "POST":
            try:
                file = request.FILES['file']
            except:
                file = ""
            text = request.POST['text']
            post_type = request.POST['postType']
            client_type = request.POST['type']
            post = telegramPost.objects.create(
                file=file,
                text=text,
                post_type=post_type,
                client_type=client_type
                )
                
        #     messages.success(request, "Post qo'shildi")
        #     return redirect('telegram')
        
        # if 'post' in request.GET:
            # post = telegramPost.objects.get(id=request.GET['post'])
            
            # subs = Client.objects.filter(~Q(telegram=""))
            # subs = subs.filter(~Q(telegram__isnull=True))
            
            if post.client_type != 'all':
                subs = Client.objects.filter(type=post.client_type).exclude(bot_user=None)
            else:
                subs = Bot_user.objects.exclude(phone = None)
            # from data.config import BOT_TOKEN
            # bot = telegram.Bot(token = BOT_TOKEN)
            # for sub in subs:
            #     if post.client_type == 'all':
            #         user_id = sub.user_id
            #     else:
            #         user_id = sub.bot_user.user_id
            #     if post.post_type == 'text':
            #         bot.sendMessage(chat_id=user_id, text = post.text, parse_mode=telegram.ParseMode.HTML)
            #     elif post.post_type == 'photo':
            #         bot.sendPhoto(chat_id=user_id, photo=post.file, caption=post.text)
            #     elif post.post_type == 'video':
            #         bot.sendVideo(chat_id=user_id, video=post.file, caption=post.text)
            #     elif post.post_type == 'audio':
            #         bot.sendAudio(chat_id=user_id, audio=post.file, caption=post.text)
            #     elif post.post_type == 'document':
            #         bot.sendDocument(chat_id=user_id, document=post.file, caption=post.text)
            #     else:
            #         messages.error(request, "Xatolik( Qaytadan boshlang")
            #         return redirect("telegram")
            messages.success(request, f"Xabar {subs.count()}ta obunachiga yuborildi")

        return render(request, 'base/telegram.html', {'profile':profile, 'posts':posts, 'telegrams':teleusers, 'bot_users': bot_users})
    else:
        return render(request, 'error-404.html')

class BotuserEditView(LoginRequiredMixin, UpdateView):
    model = Bot_user
    form_class = Bot_userForm
    template_name = 'base/forms.html'
    success_url = '/telegram'

@login_required
def calendar(request):
    all_list = {'Январь': [], 'Февраль': [], 'Март': [], 'Апрель': [], 'Май': [], 'Июнь': [], 'Июль': [], 'Август': [], 'Сентябрь': [], 'Октябрь': [], 'Ноябрь': [], 'Декабрь': []}
    n_month = 1
    days_of_the_months = {'Январь': 31, 'Февраль': 28, 'Март': 31, 'Апрель': 30, 'Май': 31, 'Июнь': 30, 'Июль': 31, 'Август': 31, 'Сентябрь':30, 'Октябрь': 31, 'Ноябрь': 30, 'Декабрь': 31}
    if int(datetime.datetime.today().year) % 4 == 0:
        days_of_the_months['Февраль'] = 29
    for i in days_of_the_months:
        for d in range(1, days_of_the_months[i]+1):
            # d = str(d)
            # if len(d) == 1:
            #     d = '0'+d
            
            # n_month = str(n_month_)
            # if len(n_month) == 1:
            #     n_month = '0'+n_month
            
            # obj = Client.objects.filter(date_birthday__icontains='-{}-'.format(str(n_month))).filter(date_birthday__endswith='-'+str(d))
            obj = Notes.objects.filter(Q(period__month=n_month) & Q(period__day=d)).filter(status='0')
            if obj:
                ps = ''
                for note in obj:
                    if note.client:
                        ps += note.client.name + ',\n'
                all_list[i].append(ps)
            else:
                all_list[i].append('None')


        
        n_month+=1
    for i in all_list:
        for p in range(5):
            all_list[i].append('Null')
    context = {'all_list': all_list}
    return render(request, 'base/calendar.html', context)

@login_required
def get_auction_info_file(request, client_pk, type):
    today = date.today()
    day = today.day
    month = services.get_uzb_month(today.month)
    year = today.year
    client = Client.objects.get(pk=client_pk)
    # p = os.listdir(os.path.join(BASE_DIR, 'files/main'))
    document = Document(os.path.join(BASE_DIR, 'files/main/auction.docx'))
    changes_in_text = [
        ('No', client.id), ('day', day), ('month', month), ('year', year), 
        ('name', client.name), ('phone1', client.phone1), ('phone2', client.phone2 or ''), 
        ('jshshir', client.jshshir), ('address2', client.address2), ('address1', client.address), 
        ('start_price', nf(client.start_price)), ('end_price', nf(client.end_price)), ('up_to_price', nf(client.up_to_price)), 
        ('pledge', nf(client.pledge)), ('overall_price', nf(client.overall_price)), ('win_value', nf(client.win_value)), 
        ('service_fee', nf(client.service_fee)), ('overall_payment', nf(client.overall_payment)) 
        ]
    try:
        for paragraph in document.paragraphs:
            font = paragraph.style.font
            font.size = Pt(14)
            if paragraph.text == 'E-IJRO AUKSION':
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for text in changes_in_text:
                paragraph.text = paragraph.text.replace('{'+text[0]+'}', str(text[1]))
    except Exception as exp:
        print(exp)
    document.save('static/auction/{}.docx'.format(client.pk))
    p = os.path.abspath('static/auction/{}.docx'.format(client.pk))
    if type == 'word':
        f = open(p, 'rb')
    elif type == 'pdf':
        os.system('unoconv -f pdf {}'.format(p))
        f = open('static/auction/{}.pdf'.format(client.pk), 'rb')
    return FileResponse(f)

@login_required
def get_carnumber_info_file(request, client_pk, type):
    today = date.today()
    day = today.day
    month = services.get_uzb_month(today.month)
    year = today.year
    client = Client.objects.get(pk=client_pk)
    # p = os.listdir(os.path.join(BASE_DIR, 'files/main'))
    document = Document(os.path.join(BASE_DIR, 'files/main/carnumber.docx'))
    changes_in_text = [
        ('No', client.id), ('day', day), ('month', month), ('year', year), 
        ('name', client.name), ('owner', client.owner), ('phone1', client.phone1), ('phone2', client.phone2 or ''), 
        ('jshshir', client.jshshir), ('address2', client.address2), ('address', client.address), 
        ('start_price', nf(client.start_price)), ('sold_price', nf(client.sold_price)), 
        ('stock_market_price', nf(client.stock_market_price)), ('pledge', nf(client.pledge)), 
        ('overall_price', nf(client.overall_price)), ('win_value', nf(client.win_value)), 
        ('service_fee', nf(client.service_fee)), ('overall_payment', nf(client.overall_payment)) 
        ]
    try:
        for paragraph in document.paragraphs:
            font = paragraph.style.font
            font.size = Pt(14)
            for text in changes_in_text:
                paragraph.text = paragraph.text.replace('{'+text[0]+'}', str(text[1]))
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.text = para.text.replace('{order1}', car_number(client.order1) or '')
                        para.text = para.text.replace('{order2}', car_number(client.order2) or '')
                        para.text = para.text.replace('{order3}', car_number(client.order3) or '')

    except Exception as exp:
        print(exp)
    document.save('static/auction/{}.docx'.format(client.pk))
    p = os.path.abspath('static/auction/{}.docx'.format(client.pk))
    if type == 'word':
        f = open(p, 'rb')
    elif type == 'pdf':
        os.system('unoconv -f pdf {}'.format(p))
        f = open('static/auction/{}.pdf'.format(client.pk), 'rb')
    return FileResponse(f)


@login_required
def templates(request):
    user = request.user
    profile = Profile.objects.get(user=user)
    list = Template.objects.all()
    context = {'templates': list, 'profile': profile}
    return render(request, 'base/templates.html', context)

class TemplateEditView(LoginRequiredMixin, UpdateView):
    model = Template
    form_class = TemplateForm
    template_name = 'base/forms.html'
    success_url = '/templates'
    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)
    
    def get_context_data(self, **kwargs):
        profile = Profile.objects.get(user=self.request.user)
        context = super().get_context_data(**kwargs)
        context['profile'] = profile
        context['view'] = 'template'
        return context

class TemplateCreateView(LoginRequiredMixin, CreateView):
    model = Template
    form_class = TemplateForm
    template_name = 'base/forms.html'
    success_url = '/templates'
    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)
    
    def get_context_data(self, **kwargs):
        profile = Profile.objects.get(user=self.request.user)
        context = super().get_context_data(**kwargs)
        context['profile'] = profile
        context['view'] = 'create-template'
        return context


@login_required
def template_delete(request, pk):
    template = Template.objects.get(pk=pk)
    template.delete()
    return redirect(templates)

@login_required
def keys(request, active_type=None):
    active_type = None if active_type == 'None' else active_type
    profile = Profile.objects.get(user=request.user)
    if request.method == 'POST':
        files = request.FILES
        success_messages = ''
        error_messages = ''

        for file in request.FILES.getlist('files'):
            *args, file_format = str(file).split('.')
            if file_format == 'pfx':
                file_path = handle_uploaded_file(file)
                f = open(file_path, 'rb')
                symbols = "QWERTYUIOPLKJHGFDSAZXCVBNMqwertyuiopasdfghjklmnbvcxz.,';:-=+1234567890"
                all_text = f.read()
                text = ''
                for i in all_text:
                    i = chr(i)
                    try:
                        if str(i) in symbols:
                            text += str(i)
                    except:
                        None
                try:
                    info_list = text.split(',')
                    for i in text.split(','):
                        if not 'name=' in i:
                            info_list.remove(i)
                        else:
                            break
                    
                    # print(info_list)

                    if 'businesscategory=' in info_list[10]: # Yuridk
                        type_ = 'yuridik'

                    elif 'o==' in info_list[5]: # Jismoniy
                        type_ = 'jismoniy'
   
                    elif 'o=' in info_list[5]: # YaTT
                        type_ = 'ytt'

                    else:
                        type_ = None
                    
                    name = info_list[0].split('=')[1]
                    surname = info_list[1].split('=')[1]
                    uid = None
                    for i in info_list[2:20]:
                        if 'validto=' in i:
                            validto = i.split('=')[1][:10]
                        elif '1.2.860.3.16.1.2=' in i:
                            jshshir = i.split('=')[1]
                        elif 'uid=' in i:
                            uid = i.split('=')[1]
                    if not uid:
                        uid = jshshir

                    year, month, day = map(int, validto.split('.'))
                    print(type_, name, surname, validto, jshshir, uid) # check all values are available

                    if not Key.objects.filter(jshshir=jshshir, key_exp__year=year, key_exp__month=month, key_exp__day=day):
                        Key.objects.create(
                            added_by = profile,
                            name = surname + ' ' + name,
                            jshshir = jshshir,
                            inn = uid,
                            key = file_path,
                            key_exp = date(year, month, day),
                            type = type_
                        )
                        success_messages += '+'
                    else:
                        error_messages += "Bu kalit allaqachon yuklangan {}\n".format(file)

                except:
                    error_messages += "Faylni o'qishda xatolik {}\n".format(file)
            
            else:
                error_messages += "Noto'g'ri fayl {}\n".format(file)
    
        # messages
        if success_messages:
            messages.success(request, "Kalitlar muvaffaqiyatli qo'shildi")
        if error_messages:
            messages.error(request, error_messages)

    query = Key.objects.all().order_by('key_exp')
    
    # determine types
    if active_type == 'all':
        types = [('all', 'Barcha')]
        type_links = list(Key.TYPE_CHOICES)
        type_links.insert(0, (None, 'Tur kiritilmagan'))
    else:
        types = list(Key.TYPE_CHOICES)
        types.insert(0, (None, 'Tur kiritilmagan'))
        type_links = [('all', 'Barcha')]
    
    context = {'keys': query, 'types': types, 'type_links': type_links, 'active_type': active_type, 'profile': profile}
    return render(request, 'base/keys.html', context)


@login_required
@permission_required('base.change_key')
def change_key_type(request, key_pk, type, active_type):
    key_obj = Key.objects.get(pk=key_pk)
    key_obj.type = type
    key_obj.save()
    return redirect(keys, active_type=active_type)


@csrf_exempt
def bot_webhook(request):

    if ENVIRONMENT == 'local':
        updater.start_polling()
    else:
        update = Update.de_json(json.loads(request.body.decode('utf-8')), dp.bot)
        dp.process_update(update)
    return HttpResponse('Bot started!')

def get_file(request, folder, file):
    f = open('static/{}/{}'.format(folder, file), 'rb')
    return FileResponse(f)

def get_template(request, folder, file):
    f = open('static/templates/{}/{}'.format(folder, file), 'rb')
    return FileResponse(f)

