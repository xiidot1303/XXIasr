from base.views import *
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

@login_required(login_url='login')
def decree_list(request):
    profile = Profile.objects.get(user = request.user)
    # get decrees which is submitted by me | get all decress to admins
    if profile.status == 'admin' or profile.status == 'superadmin':
        submitted_decrees = Decree.objects.filter().exclude(status = 'done')
    else:
        submitted_decrees = Decree.objects.filter(submitter=profile).exclude(status = 'done')
    
    # get decrees which is receiver is me
    received_decrees = Decree.objects.filter(receiver=profile).exclude(status = 'done')

    context = {
        'submitted_decrees': submitted_decrees, 
        'received_decrees': received_decrees,
        'profile': profile,
        }

    return render(request, 'decree/decree_list.html', context)

class DecreeCreateView(LoginRequiredMixin, PermissionRequiredMixin, CreateView):
    model = Decree
    form_class = DecreeForm
    template_name = 'base/forms.html'
    success_url = '/decree-list'
    login_url = 'login'
    permission_required = 'base.add_decree'

    def form_valid(self, form):
        # get current profile
        profile = Profile.objects.get(user = self.request.user)
        form.instance.submitter = profile
        return super().form_valid(form)
    
    def get_context_data(self, **kwargs):
        profile = Profile.objects.get(user=self.request.user)
        context = super().get_context_data(**kwargs)
        context['profile'] = profile
        context['view'] = 'Buyruq yaratish'
        return context

class DecreeEditView(LoginRequiredMixin, PermissionRequiredMixin, UpdateView):
    model = Decree
    form_class = DecreeForm
    template_name = 'base/forms.html'
    success_url = '/decree-list'
    login_url = 'login'
    permission_required = 'base.change_decree'
    
    def form_valid(self, form):
        return super().form_valid(form)
    
    def get_context_data(self, **kwargs):
        profile = Profile.objects.get(user=self.request.user)
        context = super().get_context_data(**kwargs)
        context['profile'] = profile
        context['view'] = 'Buyruq'
        return context

class DecreeCompleteView(LoginRequiredMixin, UpdateView):
    model = Decree
    form_class = DecreeCompleteForm
    template_name = 'base/forms.html'
    success_url = '/decree-list'
    login_url = 'login'
    
    def form_valid(self, form):
        form.instance.status = 'checking'
        return super().form_valid(form)
    
    def get_context_data(self, **kwargs):
        profile = Profile.objects.get(user=self.request.user)
        context = super().get_context_data(**kwargs)
        context['profile'] = profile
        context['view'] = 'Buyruqni yakunlash'
        return context

@login_required(login_url='login')
def decree_receive(request, pk):
    decree = Decree.objects.get(pk=pk)
    decree.status = 'received'
    decree.save()
    return redirect('decree_list')


@login_required(login_url='login')
@permission_required('base.change_decree')
def decree_check(request, pk, status):
    decree = Decree.objects.get(pk=pk)
    decree.status = status
    decree.save()
    return redirect('decree_list')


@login_required(login_url='login')
def decree_download_substance(request, pk):
    decree = Decree.objects.get(pk=pk)
    substance = decree.substance

    # Create a new Document
    doc = Document()

    # Add some text to the document
    heading = f"№{decree.id} / {decree.date.strftime('%d.%m')}\n" \
        f"{decree.date.strftime('%d.%m.%Y')} y."
    doc.add_heading(heading, level=1)
    # Add two empty paragraphs to create space
    doc.add_paragraph('')
    doc.add_paragraph('')
    # Add the text "qaror" in the center of the document with a larger font size
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Qaror')
    run.font.size = Pt(20)  # Set font size
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the paragraph

    doc.add_paragraph(substance)

    # Save the document to a BytesIO object
    from io import BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Create the HTTP response
    response = HttpResponse(file_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Buyruq mazmuni ({decree.receiver.name}).docx"'

    return response